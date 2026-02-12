from pathlib import Path
import base64
import textwrap
import pandas as pd
import numpy as np
import streamlit as st
import matplotlib.pyplot as plt
from urllib.parse import quote_plus
import os
from html import escape
import time
import requests


try:
    from docx import Document
except Exception:
    Document = None

st.set_page_config(page_title="Markmentum | Morning Compass", layout="wide")

from utils.auth import restore_auth_from_cookie, set_auth_cookie



VERIFY_URL = "https://admin.memberstack.com/members/verify-token"

def verify_memberstack_token(token: str) -> dict | None:
    secret = os.environ.get("MEMBERSTACK_SECRET_KEY")
    if not secret:
        return None

    try:
        r = requests.post(
            VERIFY_URL,
            headers={
                "X-API-KEY": secret,
                "Content-Type": "application/json",
            },
            json={"token": token},
            timeout=10,
        )
        if r.status_code != 200:
            return None

        return r.json().get("data")  # includes id/type/iat/exp/aud/iss
    except Exception:
        return None


def establish_auth() -> bool:
    # Grab once
    ms_token = st.query_params.get("ms_session")

    # If token is in URL, stash it (once) and FORCE a clean reload (no rerun)
    if ms_token:
        if not st.session_state.get("_pending_ms_session"):
            st.session_state["_pending_ms_session"] = ms_token

        # Hard reload to clean URL (guarantees token disappears from address bar)
        st.markdown(
            """<meta http-equiv="refresh" content="0; url=https://app.markmentumresearch.com/" />""",
            unsafe_allow_html=True,
        )
        st.stop()

    # If already authenticated, we're good (also clear any leftover pending stash)
    if st.session_state.get("authenticated") is True:
        st.session_state.pop("_pending_ms_session", None)
        return True

    # Grab token from pending stash (URL token would have been stashed + removed above)
    token = st.session_state.get("_pending_ms_session")

    # Now we’re on a clean URL; proceed using the stashed token
    if token:
        verified = verify_memberstack_token(token)
        if not verified:
            st.session_state["authenticated"] = False
            st.session_state.pop("_pending_ms_session", None)
            return False

        expected_aud = os.environ.get("MEMBERSTACK_APP_ID")
        if expected_aud and verified.get("aud") != expected_aud:
            st.session_state["authenticated"] = False
            st.session_state.pop("_pending_ms_session", None)
            return False

        now = int(time.time())
        if isinstance(verified.get("exp"), int) and verified["exp"] < now:
            st.session_state["authenticated"] = False
            st.session_state.pop("_pending_ms_session", None)
            return False

        member_id = (verified.get("id") or "").strip()
        st.session_state["authenticated"] = True
        st.session_state["member_id"] = member_id
        st.session_state["auth_checked_at"] = now

        # Write cookie AFTER URL is already clean
        if member_id:
            set_auth_cookie(member_id)
            st.rerun()

        # Done with the stashed token
        st.session_state.pop("_pending_ms_session", None)

        return True


    # No token -> try cookie restore
    if restore_auth_from_cookie():
        st.session_state["authenticated"] = True
        return True

    st.session_state["authenticated"] = False
    return False

# --- Gate Morning Compass ---
if not establish_auth():
    home_url = "https://www.markmentumresearch.com"
    st.markdown(
        f"""<meta http-equiv="refresh" content="0; url={home_url}" />""",
        unsafe_allow_html=True,
    )
    st.stop()

# -------------------------
# Page & shared style
# -------------------------

st.cache_data.clear()
# ---- LAYOUT & WIDTH TUNING (Cloud parity + your constraints) ----





# -------------------------
# Paths (portable for Cloud)
# -------------------------
_here = Path(__file__).resolve().parent
APP_DIR = _here if _here.name != "pages" else _here.parent

DATA_DIR   = APP_DIR / "data"
ASSETS_DIR = APP_DIR / "assets"
LOGO_PATH  = ASSETS_DIR / "markmentum_logo.png"


# -------------------------
# Header (logo centered)
# -------------------------
def _image_to_base64(path: Path) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()


if LOGO_PATH.exists():
    st.markdown(
        f"""
        <div style="text-align:center; margin: 8px 0 16px;">
            <img src="data:image/png;base64,{_image_to_base64(LOGO_PATH)}" width="440">
        </div>
        """,
        unsafe_allow_html=True,
    )


# -------------------------
# Helpers
# -------------------------
ADV_VALUE_KEY  = "dd_show_advanced_charts_value"
INFO_VALUE_KEY = "dd_show_information_charts_value"

def _mk_ticker_link(ticker: str) -> str:
    t = (ticker or "").strip().upper()
    if not t:
        return ""

    # current toggle states (default False if not set yet)
    adv_on  = st.session_state.get(ADV_VALUE_KEY, False)
    info_on = st.session_state.get(INFO_VALUE_KEY, False)

    adv_flag  = "1" if adv_on  else "0"
    info_flag = "1" if info_on else "0"

    return (
        f'<a href="?page=Deep%20Dive'
        f'&ticker={quote_plus(t)}'
        f'&adv={adv_flag}'
        f'&info={info_flag}" '
        f'target="_self" rel="noopener" '
        f'style="text-decoration:none; font-weight:600;">{t}</a>'
    )

qp = st.query_params
dest = (qp.get("page") or "").strip().lower()

if dest.replace("%20", " ") == "deep dive":
    t = (qp.get("ticker") or "").strip().upper()
    adv_qp  = (qp.get("adv")  or "0").strip()
    info_qp = (qp.get("info") or "0").strip()

    if t:
        st.session_state["ticker"] = t

        # restore toggle master values for Deep Dive
        st.session_state[ADV_VALUE_KEY]  = (adv_qp == "1")
        st.session_state[INFO_VALUE_KEY] = (info_qp == "1")

        # clean URL – we don't need adv/info/ticker in query params anymore
        st.query_params.clear()
        st.query_params["ticker"] = t

        st.switch_page("pages/08_Deep_Dive_Dashboard.py")

def row_spacer(height_px: int = 14):
    st.markdown(f"<div style='height:{height_px}px'></div>", unsafe_allow_html=True)

# ---------- Shared formatters ----------
def fmt_num(x, nd=2):
    try:
        if pd.isna(x): return ""
        return f"{float(x):,.{nd}f}"
    except Exception:
        return ""

def fmt_pct(x, nd=2):
    try:
        if pd.isna(x): return ""
        return f"{float(x)*100:,.{nd}f}%"
    except Exception:
        return ""

def fmt_int(x):
    try:
        if pd.isna(x): return ""
        return f"{int(round(float(x))):,}"
    except Exception:
        return ""

def _is_list_paragraph(paragraph) -> bool:
    try:
        return paragraph._p.pPr.numPr is not None
    except Exception:
        return False


@st.cache_data(show_spinner=False)
def load_docx_text(doc_path: str) -> str:
    """
    Reads a .docx and returns plain text (bullets preserved as '- ' lines).
    Designed for short bottom-line docs like usd_correlation_bottom_line.docx.
    """
    if Document is None:
        return "⚠️ Bottom line: python-docx is not installed (run: `pip install python-docx`)."

    if not os.path.exists(doc_path):
        return f"⚠️ Bottom line file not found: {doc_path}"

    try:
        doc = Document(doc_path)
    except Exception as e:
        return f"⚠️ Could not open bottom line file: {e}"

    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        lines.append(f"- {t}" if _is_list_paragraph(p) else t)

    return "\n".join(lines).strip()

@st.cache_data(show_spinner=False)
def load_txt_text(txt_path: str) -> str:
    if not os.path.exists(txt_path):
        return f"⚠️ Bottom line file not found: {txt_path}"
    try:
        return Path(txt_path).read_text(encoding="utf-8").strip()
    except Exception as e:
        return f"⚠️ Could not open bottom line file: {e}"

# ---------- UI renderers ----------
def mm_badge_html(x):
        try:
            if pd.isna(x):
                return ""
            v = float(x)
        except Exception:
            return ""

        if v <= -100:
            bg, alpha = "rgba(185,28,28,0.35)", 0.35   # deep red
        elif v < -25:
            bg, alpha = "rgba(239,68,68,0.28)", 0.28   # red
        elif v <= 25:
            bg, alpha = "rgba(229,231,235,1.00)", 1.00 # gray pill
        elif v < 100:
            bg, alpha = "rgba(16,185,129,0.28)", 0.28  # green
        else:
            bg, alpha = "rgba(6,95,70,0.35)", 0.35     # dark green

        label = f"{int(round(v)):,}"
        # block so it fills the cell nicely; cell stays right-aligned from CSS
        return f'<span style="display:block; background:{bg}; padding:0 4px; border-radius:2px;">{label}</span>'


def rr_tinted_html(x, cap=3.0):
        try:
            if pd.isna(x): 
                return ""
            v = float(x)
        except Exception:
            return ""

        # scale 0..1 (capped), keep near-zero very light
        s = min(abs(v) / cap, 1.0)
        alpha = 0.12 + 0.28 * s     # 0.12 → 0.40 opacity

        if v > 0:
            # green (tailwind-ish 10B981)
            bg = f"rgba(16,185,129,{alpha:.3f})"
        elif v < 0:
            # red (EF4444)
            bg = f"rgba(239,68,68,{alpha:.3f})"
        else:
            bg = "transparent"

        # numeric label with 1 decimal, same as before
        label = f"{v:,.1f}"
        return f'<span style="display:block; background:{bg}; padding:0 4px; border-radius:2px;">{label}</span>'

# =========================
# Timeframe config
# =========================
TIMEFRAMES = {
    "Daily": {
        "ids": {"main": 73, "leaders": 74, "mm": 75, "category": 76, "delta": 77},
        "cols": {"ret": "daily_Return", "pr_low": "day_pr_low", "pr_high": "day_pr_high", "rr": "day_rr_ratio"},
        "txt": "bottom_line_daily.txt",
        "card_title": "Daily Macro Orientation",
        "card_title2":"Daily Top Five Leaders/Laggards by % Change",
        "card_title3":"Daily Top Five Leaders/Laggards by MM Score",
        "card_title4":"Daily Top Five Leaders/Laggards by MM Score Change",
        "card_title5":"Daily Category Snapshot",
    },
    "Weekly": {
        "ids": {"main": 78, "leaders": 79, "mm": 80, "category": 81, "delta": 82},
        "cols": {"ret": "weekly_Return", "pr_low": "week_pr_low", "pr_high": "week_pr_high", "rr": "week_rr_ratio"},
        "txt": "bottom_line_weekly.txt",
        "card_title": "Weekly Macro Orientation",
        "card_title2":"Weekly Top Five Leaders/Laggards by % Change",
        "card_title3":"Weekly Top Five Leaders/Laggards by MM Score",
        "card_title4":"Weekly Top Five Leaders/Laggards by MM Score Change",
        "card_title5":"Weekly Category Snapshot",
         
    },
    "Monthly": {
        "ids": {"main": 83, "leaders": 84, "mm": 85, "category": 86, "delta": 87},
        "cols": {"ret": "monthly_Return", "pr_low": "month_pr_low", "pr_high": "month_pr_high", "rr": "month_rr_ratio"},
        "txt": "bottom_line_monthly.txt",
        "card_title": "Monthly Macro Orientation",
        "card_title2":"Monthly Top Five Leaders/Laggards by % Change",
        "card_title3":"Monthly Top Five Leaders/Laggards by MM Score",
        "card_title4":"Monthly Top Five Leaders/Laggards by MM Score Change",
        "card_title5":"Monthly Category Snapshot",
         
    },
}

# Centered timeframe selector under the page title (we’ll render title after we know the date)
def timeframe_selector(default="Daily"):
    c1, c2, c3 = st.columns([1, 0.8, 1])
    with c2:
        return st.selectbox("Timeframe", list(TIMEFRAMES.keys()), index=list(TIMEFRAMES.keys()).index(default), label_visibility="collapsed")

@st.cache_data(show_spinner=False)
def load_csv_by_id(n: int, base_dir: Path) -> pd.DataFrame:
    p = base_dir / f"qry_graph_data_{n}.csv"
    if not p.exists():
        return pd.DataFrame()
    return pd.read_csv(p)



# -------------------------
# Morning Compass – styling
# -------------------------
# -------------------------
# Morning Compass – styling
# -------------------------
st.markdown("""
<style>
/* Center the single card on the page */
.card-wrap { display:flex; justify-content:center; }
.card { 
  border:1px solid #cfcfcf; border-radius:8px; background:#fff;
  padding:12px 12px 8px 12px; width:100%;
  max-width:1320px;  /* was 1120px -> more room so names show */
}

/* Table styling to match Daily Overview */
.tbl { border-collapse: collapse; width: 100%; table-layout: fixed; }
.tbl th, .tbl td {
  border:1px solid #d9d9d9; padding:6px 8px; font-size:13px;
  overflow:hidden; text-overflow:ellipsis;
}
.tbl th { background:#f2f2f2; font-weight:700; color:#1a1a1a; text-align:left; }

/* Alignment rules */
.tbl th:nth-child(2), .tbl td:nth-child(2) { text-align:center; }    /* Ticker centered */

/* HEADERS from Close..MM Score Change centered */
.tbl th:nth-child(n+3) { text-align:center; }

/* CELLS from Close..MM Score Change right-aligned */
.tbl td:nth-child(n+3) { text-align:right; white-space:nowrap; }

/* Name column = 40ch, allow wrapping so full name shows */
.tbl col.col-name { min-width:40ch; width:40ch; max-width:40ch; }
.tbl th:nth-child(1), .tbl td:nth-child(1) {
  white-space:normal;               /* allow wrap */
  overflow:visible; text-overflow:clip;
}

/* Keep ticker links bold without underline */
.tbl a { text-decoration:none; font-weight:600; }
            
/* Correlation tables: center 15D/30D/90D cells */
.tbl.corr th:nth-child(n+2),
.tbl.corr td:nth-child(n+2) {
  text-align: center !important;
}

/* keep correlation numbers on one line */
.tbl.corr td:nth-child(n+2) { white-space: nowrap; }


</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>
/* bottom line inside the card, attached to table border */
.bl {
  border-top: 1px solid #e5e5e5;
  margin-top: 8px;
  padding-top: 10px;
  font-size: 13px;
  line-height: 1.45;
  color: #1a1a1a;
}
.card .note {
  font-size: 0.85em;
  color: #6c757d;   /* muted gray */
  line-height: 1.3;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>
/* ========== Bigger, easier-to-grab scrollbars (light theme) ========== */

/* Global page scrollbar — Chromium/WebKit */
:root::-webkit-scrollbar        { width: 16px; height: 16px; }
:root::-webkit-scrollbar-track  { background: #f2f2f2; }
:root::-webkit-scrollbar-thumb  { background: #bdbdbd; border-radius: 8px; border: 3px solid #f2f2f2; }
:root::-webkit-scrollbar-thumb:hover { background: #9a9a9a; }

/* Inner scrollables (tables/dataframes/expander bodies) — Chromium/WebKit */
div[tabindex="0"]::-webkit-scrollbar        { width: 14px; height: 14px; }
div[tabindex="0"]::-webkit-scrollbar-track  { background: #f2f2f2; }
div[tabindex="0"]::-webkit-scrollbar-thumb  { background: #bdbdbd; border-radius: 8px; border: 3px solid #f2f2f2; }
div[tabindex="0"]::-webkit-scrollbar-thumb:hover { background: #9a9a9a; }

/* Firefox */
html { scrollbar-width: thick; scrollbar-color: #bdbdbd #f2f2f2; }
</style>
""", unsafe_allow_html=True)

# -------------------------
# Load Morning Compass CSV
# -------------------------
# -------------------------
# Select timeframe and load MAIN csv for date/title
# -------------------------
sel_tf = timeframe_selector(default="Daily")
cfg_tf  = TIMEFRAMES[sel_tf]

df_main = load_csv_by_id(cfg_tf["ids"]["main"], DATA_DIR)

# Page title under the logo (date pulled from selected timeframe's main csv)
date_str = ""
if not df_main.empty and "Date" in df_main.columns:
    asof = pd.to_datetime(df_main["Date"], errors="coerce").max()
    if pd.notna(asof):
        date_str = f"{asof.month}/{asof.day}/{asof.year}"

st.markdown(
    f"""
    <div style="text-align:center; margin:-6px 0 8px;
                font-size:18px; font-weight:600; color:#1a1a1a;">
        Morning Compass – {date_str}
    </div>
    """,
    unsafe_allow_html=True,
)
row_spacer(6)


def render_correlation_card(title: str, csv_id: int, docx_name: str):
    df = load_csv_by_id(csv_id, DATA_DIR)
    if df.empty:
        st.info(f"{title}: `qry_graph_data_{csv_id}.csv` not found.")
        return

    # format numeric columns
    df_fmt = df.copy()
    for c in ["15D", "30D", "90D"]:
        if c in df_fmt.columns:
            df_fmt[c] = df_fmt[c].map(lambda v: fmt_num(v, 2))

    table_html = df_fmt.to_html(index=False, classes="tbl corr", escape=False, border=0)
    table_html = table_html.replace('class="dataframe tbl corr"', 'class="tbl corr"')

    # load bottom line from docx (plain text)
    docx_path = (DATA_DIR / docx_name).resolve()
    bl_text = load_docx_text(str(docx_path))
    bl_html_safe = escape(bl_text).replace("\n", "<br>")
    if csv_id == 93:
        note_text = (
            "Note: USD correlations use the U.S. Dollar Index (DXY), a trade-weighted FX index. "
            "15D/30D/90D are trading-day windows. Correlation ranges from -1 to +1. "
            "Negative = tends to move opposite. Positive = tends to move together."
        )
    else:
        note_text = (
            "Note: Rate correlations use the 10-Year U.S. Treasury yield (TNX) as the rates proxy. "
            "15D/30D/90D are trading-day windows. Correlation ranges from -1 to +1. "
            "Negative = tends to move opposite. Positive = tends to move together."
        )

    note_html_safe = escape(note_text)

    card_html = f"""
    <div class="card-wrap">
      <div class="card">
        <h3 style="margin:0 0 8px 0; font-size:16px; font-weight:700; color:#1a1a1a;">
          {title}
        </h3>
        {table_html}
        <div class="bl">{bl_html_safe}</div>
        <div class="bl note">{note_html_safe}</div>
      </div>
    </div>
    """
    st.markdown(card_html, unsafe_allow_html=True)


# -------------------------
# Card 1: Morning Compass table (uses selected timeframe)
# -------------------------
cols = cfg_tf["cols"]
req = ["Date","Ticker","Ticker_name","Close",
       cols["ret"], cols["pr_low"], cols["pr_high"],
       cols["rr"], "model_score","model_score_delta"]

if df_main.empty or not all(c in df_main.columns for c in req):
    st.info(f"Morning Compass: `qry_graph_data_{cfg_tf['ids']['main']}.csv` is missing or columns are incomplete.")
else:
    df_render = df_main.copy()
    df_render["Ticker"] = df_render["Ticker"].apply(_mk_ticker_link)

    # formatters (reuse existing helpers defined above)
    def fmt_num(x, nd=2):
        try:
            if pd.isna(x): return ""
            return f"{float(x):,.{nd}f}"
        except Exception: return ""
    def fmt_pct(x, nd=2):
        try:
            if pd.isna(x): return ""
            return f"{float(x)*100:,.{nd}f}%"
        except Exception: return ""
    def fmt_int(x):
        try:
            if pd.isna(x): return ""
            return f"{int(round(float(x))):,}"
        except Exception: return ""

    df_card = pd.DataFrame({
        "Name":           df_render["Ticker_name"],
        "Ticker":         df_render["Ticker"],
        "Close":          df_render["Close"].map(lambda v: fmt_num(v, 2)),
        "% Change":       df_render[cols["ret"]].map(lambda v: fmt_pct(v, 2)),
        "Probable Low":   df_render[cols["pr_low"]].map(lambda v: fmt_num(v, 2)),
        "Probable High":  df_render[cols["pr_high"]].map(lambda v: fmt_num(v, 2)),
        "Risk / Reward":  df_render[cols["rr"]].map(rr_tinted_html),
        "MM Score":       df_render["model_score"].map(mm_badge_html),
        "Δ MM Score":df_render["model_score_delta"].map(fmt_int),
    })

    table_html = df_card.to_html(index=False, classes="tbl", escape=False, border=0)
    table_html = table_html.replace('class="dataframe tbl"', 'class="tbl"')
    colgroup = """
    <colgroup>
      <col class="col-name">
      <col>
      <col>
      <col>
      <col>
      <col>
      <col>
      <col>
      <col>
    </colgroup>
    """.strip()
    table_html = table_html.replace('<table class="tbl">', f'<table class="tbl">{colgroup}', 1)

    # Bottom line docx switches with timeframe
    import os
    try:
        from docx import Document
    except Exception:
        Document = None

    def _is_list_paragraph(paragraph) -> bool:
        try:
            return paragraph._p.pPr.numPr is not None
        except Exception:
            return False

    @st.cache_data(show_spinner=False)
    def load_market_read_md(doc_path: str) -> str:
        if Document is None:
            return "⚠️ **Market Read**: python-docx is not installed (run: `pip install python-docx`)."
        if not os.path.exists(doc_path):
            return f"⚠️ **Market Read** file not found: `{doc_path}`"
        try:
            doc = Document(doc_path)
        except Exception as e:
            return f"⚠️ Could not open **Market Read** file `{doc_path}`: {e}"
        lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            lines.append(f"- {text}" if _is_list_paragraph(p) else text)
        for i, l in enumerate(lines):
            if l.startswith("Market Read:") and "The model is saying:" in l:
                left, right = l.split("The model is saying:", 1)
                lines[i] = left.strip()
                lines.insert(i + 1, "The model is saying:")
                if right.strip():
                    lines.insert(i + 2, right.strip())
                break
        return "\n\n".join(lines)

    from html import escape
    txt_path = (DATA_DIR / cfg_tf["txt"]).resolve()
    bl_text = load_txt_text(str(txt_path)).strip()
    bl_html_safe = escape(bl_text).replace("\n", "<br>")
    note_text = "Note: MM Score → Rules-based contrarian score designed to avoid chasing stretch, identify crowding, and size conviction sensibly."
    note_html_safe = escape(note_text)

    card_html = f'''
    <div class="card-wrap">
      <div class="card">
        <h3 style="margin:0 0 8px 0; font-size:16px; font-weight:700; color:#1a1a1a;">
          {cfg_tf["card_title"]}
        </h3>
        {table_html}
        <div class="bl">{bl_html_safe}</div>
        <div class="bl note">{note_html_safe}</div>
      </div>
    </div>
    '''
    st.markdown(card_html, unsafe_allow_html=True)

row_spacer(10)
# =========================
# USD & Rates Correlations (Daily only)
# =========================
if sel_tf == "Daily":
    render_correlation_card(
        title="USD Correlations",
        csv_id=93,
        docx_name="usd_correlation_bottom_line.docx",
    )

    row_spacer(10)

    render_correlation_card(
        title="Rates Correlations",
        csv_id=94,
        docx_name="tnx_correlation_bottom_line.docx",
    )

    #row_spacer(6)



# -------------------------
# Card 2: Leaders/Laggard by % Change
# -------------------------
# =========================
# Card 2: Leaders/Laggards by % Change  (Top/Bottom 5 from selected timeframe)
# =========================
df74 = load_csv_by_id(cfg_tf["ids"]["leaders"], DATA_DIR)
req74 = ["Date","Ticker","Ticker_name","Close",
         cols["ret"], cols["pr_low"], cols["pr_high"], cols["rr"], "model_score","model_score_delta"]

if df74.empty or not all(c in df74.columns for c in req74):
    row_spacer(8)
    st.info(f"Top Five Leaders/Laggards by % Change: `qry_graph_data_{cfg_tf['ids']['leaders']}.csv` is missing or columns are incomplete.")
else:
    d = df74.copy()
    d["Ticker"] = d["Ticker"].apply(_mk_ticker_link)

    df_74_card = pd.DataFrame({
        "Name":           d["Ticker_name"],
        "Ticker":         d["Ticker"],
        "Close":          d["Close"].map(lambda v: fmt_num(v, 2)),
        "% Change":       d[cols["ret"]].map(lambda v: fmt_pct(v, 2)),
        "Probable Low":   d[cols["pr_low"]].map(lambda v: fmt_num(v, 2)),
        "Probable High":  d[cols["pr_high"]].map(lambda v: fmt_num(v, 2)),
        "Risk / Reward":  d[cols["rr"]].map(rr_tinted_html),
        "MM Score":       d["model_score"].map(mm_badge_html),
        "Δ MM Score":d["model_score_delta"].map(fmt_int),
    })

    tbl_html_74 = df_74_card.to_html(index=False, classes="tbl", escape=False, border=0)
    tbl_html_74 = tbl_html_74.replace('class="dataframe tbl"', 'class="tbl"')
    colgroup = """
    <colgroup>
      <col class="col-name"><col><col><col><col><col><col><col><col>
    </colgroup>
    """.strip()
    tbl_html_74 = tbl_html_74.replace('<table class="tbl">', f'<table class="tbl">{colgroup}', 1)

    note_text = "Note: MM Score → Rules-based contrarian score designed to avoid chasing stretch, identify crowding, and size conviction sensibly."
    note_html_safe = escape(note_text)

    row_spacer(10)
    st.markdown(
        f"""
        <div class="card-wrap">
          <div class="card">
            <h3 style="margin:0 0 8px 0; font-size:16px; font-weight:700; color:#1a1a1a;">
              {cfg_tf["card_title2"]}
            </h3>
            {tbl_html_74}
            <div class="bl note">{note_html_safe}</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# =========================
# Card 3: Leaders/Laggards by MM Score (Top/Bottom 5)
# =========================
df75 = load_csv_by_id(cfg_tf["ids"]["mm"], DATA_DIR)
req75 = ["Date","Ticker","Ticker_name","Close",
         cols["ret"], cols["pr_low"], cols["pr_high"], cols["rr"], "model_score","model_score_delta"]

if df75.empty or not all(c in df75.columns for c in req75):
    row_spacer(8)
    st.info(f"Top Five Leaders/Laggards by MM Score: `qry_graph_data_{cfg_tf['ids']['mm']}.csv` is missing or columns are incomplete.")
else:
    d = df75.copy()
    d["Ticker"] = d["Ticker"].apply(_mk_ticker_link)

    df_75_card = pd.DataFrame({
        "Name":           d["Ticker_name"],
        "Ticker":         d["Ticker"],
        "Close":          d["Close"].map(lambda v: fmt_num(v, 2)),
        "% Change":       d[cols["ret"]].map(lambda v: fmt_pct(v, 2)),
        "Probable Low":   d[cols["pr_low"]].map(lambda v: fmt_num(v, 2)),
        "Probable High":  d[cols["pr_high"]].map(lambda v: fmt_num(v, 2)),
        "Risk / Reward":  d[cols["rr"]].map(rr_tinted_html),
        "MM Score":       d["model_score"].map(mm_badge_html),
        "Δ MM Score":d["model_score_delta"].map(fmt_int),
    })

    tbl_html_75 = df_75_card.to_html(index=False, classes="tbl", escape=False, border=0)
    tbl_html_75 = tbl_html_75.replace('class="dataframe tbl"', 'class="tbl"')
    colgroup = """
    <colgroup>
      <col class="col-name"><col><col><col><col><col><col><col><col>
    </colgroup>
    """.strip()
    tbl_html_75 = tbl_html_75.replace('<table class="tbl">', f'<table class="tbl">{colgroup}', 1)

    note_text = "Note: MM Score → Rules-based contrarian score designed to avoid chasing stretch, identify crowding, and size conviction sensibly."
    note_html_safe = escape(note_text)

    row_spacer(10)
    st.markdown(
        f"""
        <div class="card-wrap">
          <div class="card">
            <h3 style="margin:0 0 8px 0; font-size:16px; font-weight:700; color:#1a1a1a;">
              {cfg_tf["card_title3"]}
            </h3>
            {tbl_html_75}
            <div class="bl note">{note_html_safe}</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# =========================
# Card 4: Leaders/Laggards by MM Score Change (Top/Bottom 5)
# =========================
df77 = load_csv_by_id(cfg_tf["ids"]["delta"], DATA_DIR)
req77 = ["Date","Ticker","Ticker_name","Close",
         cols["ret"], cols["pr_low"], cols["pr_high"], cols["rr"], "model_score","model_score_delta"]

if df77.empty or not all(c in df77.columns for c in req77):
    row_spacer(8)
    st.info(f"Top Five Leaders/Laggards by MM Score Change: `qry_graph_data_{cfg_tf['ids']['delta']}.csv` is missing or columns are incomplete.")
else:
    d = df77.copy()
    d["Ticker"] = d["Ticker"].apply(_mk_ticker_link)

    df_77_card = pd.DataFrame({
        "Name":           d["Ticker_name"],
        "Ticker":         d["Ticker"],
        "Close":          d["Close"].map(lambda v: fmt_num(v, 2)),
        "% Change":       d[cols["ret"]].map(lambda v: fmt_pct(v, 2)),
        "Probable Low":   d[cols["pr_low"]].map(lambda v: fmt_num(v, 2)),
        "Probable High":  d[cols["pr_high"]].map(lambda v: fmt_num(v, 2)),
        "Risk / Reward":  d[cols["rr"]].map(rr_tinted_html),
        "MM Score":       d["model_score"].map(mm_badge_html),
        "Δ MM Score":d["model_score_delta"].map(fmt_int),
    })

    tbl_html_77 = df_77_card.to_html(index=False, classes="tbl", escape=False, border=0)
    tbl_html_77 = tbl_html_77.replace('class="dataframe tbl"', 'class="tbl"')
    colgroup = """
    <colgroup>
      <col class="col-name"><col><col><col><col><col><col><col><col>
    </colgroup>
    """.strip()
    tbl_html_77 = tbl_html_77.replace('<table class="tbl">', f'<table class="tbl">{colgroup}', 1)

    note_text = "Note: MM Score → Rules-based contrarian score designed to avoid chasing stretch, identify crowding, and size conviction sensibly."
    note_html_safe = escape(note_text)

    row_spacer(10)
    st.markdown(
        f"""
        <div class="card-wrap">
          <div class="card">
            <h3 style="margin:0 0 8px 0; font-size:16px; font-weight:700; color:#1a1a1a;">
              {cfg_tf["card_title4"]}
            </h3>
            {tbl_html_77}
            <div class="bl note">{note_html_safe}</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# =========================
# Card 5: (optional) Category Snapshot – uses qry_graph_data_76.csv
# =========================
# =========================
# Card 5: (optional) Category Snapshot – uses timeframe's category csv
# =========================
row_spacer(6)
show_cat = st.checkbox("View Category Snapshot", value=False)

if show_cat:
    df76 = load_csv_by_id(cfg_tf["ids"]["category"], DATA_DIR)
    req76 = ["Date","Ticker","Ticker_name","Category","Close",
             cols["ret"], cols["pr_low"], cols["pr_high"], cols["rr"], "model_score","model_score_delta"]

    if df76.empty or not all(c in df76.columns for c in req76):
        st.info(f"Category Snapshot: `qry_graph_data_{cfg_tf['ids']['category']}.csv` is missing or columns are incomplete.")
    else:
        cat_order = [
            "Sector & Style ETFs","Indices","Futures","Currencies","Commodities","Bonds","Yields","Volatility","Foreign",
            "Communication Services","Consumer Discretionary","Consumer Staples","Energy","Financials",
            "Health Care","Industrials","Information Technology","Materials","Real Estate","Utilities","MR Discretion"
        ]
        c1, c2, c3 = st.columns([1, .9, 1])
        with c2:
            present = [c for c in cat_order if c in df76["Category"].dropna().unique().tolist()]
            sel = st.selectbox("Category", present, index=0)

        d = df76[df76["Category"] == sel].copy()
        d["Ticker"] = d["Ticker"].apply(_mk_ticker_link)

        df_cat_card = pd.DataFrame({
            "Name":           d["Ticker_name"],
            "Ticker":         d["Ticker"],
            "Close":          d["Close"].map(lambda v: fmt_num(v, 2)),
            "% Change":       d[cols["ret"]].map(lambda v: fmt_pct(v, 2)),
            "Probable Low":   d[cols["pr_low"]].map(lambda v: fmt_num(v, 2)),
            "Probable High":  d[cols["pr_high"]].map(lambda v: fmt_num(v, 2)),
            "Risk / Reward":  d[cols["rr"]].map(rr_tinted_html),
            "MM Score":       d["model_score"].map(mm_badge_html),
            "Δ MM Score":d["model_score_delta"].map(fmt_int),
        })

        tbl_html_76 = df_cat_card.to_html(index=False, classes="tbl", escape=False, border=0)
        tbl_html_76 = tbl_html_76.replace('class="dataframe tbl"', 'class="tbl"')
        colgroup = """
        <colgroup>
          <col class="col-name"><col><col><col><col><col><col><col><col>
        </colgroup>
        """.strip()
        tbl_html_76 = tbl_html_76.replace('<table class="tbl">', f'<table class="tbl">{colgroup}', 1)

        note_text = "Note: MM Score → Rules-based contrarian score designed to avoid chasing stretch, identify crowding, and size conviction sensibly."
        note_html_safe = escape(note_text)

        row_spacer(6)
        st.markdown(
            f"""
            <div class="card-wrap">
              <div class="card">
                <h3 style="margin:0 0 8px 0; font-size:16px; font-weight:700; color:#1a1a1a;">
                  {cfg_tf["card_title5"]} – {sel}
                </h3>
                {tbl_html_76}
                <div class="bl note">{note_html_safe}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )



# -------------------------
# Footer disclaimer
# -------------------------
st.markdown("---")
st.markdown(
    """
    <div style="font-size: 12px; color: gray;">
    © 2026 Markmentum Research LLC. <b>Disclaimer</b>: This content is for informational purposes only. 
    Nothing herein constitutes an offer to sell, a solicitation of an offer to buy, or a recommendation regarding any security, 
    investment vehicle, or strategy. It does not represent legal, tax, accounting, or investment advice by Markmentum Research LLC 
    or its employees. The information is provided without regard to individual objectives or risk parameters and is general, 
    non-tailored, and non-specific. Sources are believed to be reliable, but accuracy and completeness are not guaranteed. 
    Markmentum Research LLC is not responsible for errors, omissions, or losses arising from use of this material. 
    Investments involve risk, and financial markets are subject to fluctuation. Consult your financial professional before 
    making investment decisions.
    </div>
    """,
    unsafe_allow_html=True,
)